from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import pandas as pd
import certifi
import requests
import tempfile
import io
from urllib.parse import quote, unquote
import os
import math
import re
from datetime import datetime

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Load the Excel file
url = 'https://raw.githubusercontent.com/trial777/combined_selected_RS/main/combined_selected_v3.xlsx'

try:
    response = requests.get(url, verify=certifi.where())
    response.raise_for_status()

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_file:
        temp_file.write(response.content)
        temp_file.seek(0)
        data = pd.read_excel(temp_file.name)

    print("Data loaded successfully!")
except Exception as e:
    print(f"Error: {e}")
    raise

# Map classified values to URL-safe identifiers
classified_url_map = {
    "Unique to Process": "unique-process",
    "Unique to People": "unique-people",
    "Unique to Technology": "unique-technology",
    "AllThreeSegments": "all-three",
    "People & Process": "people-process",
    "People & Technology": "people-technology",
    "Process & Technology": "process-technology"
}
url_to_classified = {v: k for k, v in classified_url_map.items()}

# Group data by 'Classified'
grouped_data = data.groupby("Classified")

@app.route('/<classified_url>')
def modal(classified_url):
    classified_value = url_to_classified.get(unquote(classified_url))
    if not classified_value or classified_value not in grouped_data.groups:
        return jsonify({"error": "Invalid classified value."}), 404

    filtered_data = grouped_data.get_group(classified_value)
    modal_data = filtered_data[["Reference_no", "Title", "URL"]].to_dict(orient="records")

    # Count and percentage calculations
    classified_count = len(filtered_data)
    percentage = round((classified_count / len(data)) * 100, 2)

    print(f"DEBUG: classified_url = {classified_url}")  # Add debug print

    return render_template(
        'modal.html',
        classified_url=classified_url,  # Pass this variable
        classified_value=classified_value,
        modal_data=modal_data,
        classified_count=classified_count,
        total_count=len(data),
        percentage=percentage
    )

# Ensure the required columns exist
required_columns = {"Reference_no", "Extracted Themes", "URL"}
if not required_columns.issubset(data.columns):
    raise ValueError(f"The Excel file must contain these columns: {', '.join(required_columns)}")


'''@app.route('/extracted-themes')
def extracted_themes():
    """
    Route to display paginated Reference_no and Extracted Themes.
    """
    # Drop rows where "Extracted Themes" is missing
    filtered_data = data.dropna(subset=["Extracted Themes"])

    # Pagination logic
    page = request.args.get('page', 1, type=int)  # Current page
    per_page = 25  # Number of records per page
    total_pages = math.ceil(len(filtered_data) / per_page)

    # Slice data for the current page
    start = (page - 1) * per_page
    end = start + per_page
    paginated_data = filtered_data.iloc[start:end]

    # Convert to list of dictionaries for rendering
    modal_data = paginated_data[["Reference_no", "Extracted Themes", "URL"]].to_dict(orient="records")

    return render_template(
        'extracted_themes.html',
        modal_data=modal_data,
        current_page=page,
        total_pages=total_pages
    )'''


@app.route('/themes', methods=['GET'])
def themes_page():
    """
    Route to display paginated themes with sorted reference numbers and clickable links.
    """
    # Define pagination parameters
    items_per_page = 20  # Records per page
    page = request.args.get('page', default=1, type=int)  # Default page number is 1

    # Sort data by 'Reference_no'
    sorted_data = data[['Reference_no', 'Extracted Themes', 'URL']].sort_values(by='Reference_no')

    # Pagination logic
    total_items = len(sorted_data)
    total_pages = math.ceil(total_items / items_per_page)

    # Validate the page number
    if page < 1:
        page = 1
    elif page > total_pages:
        page = total_pages

    start = (page - 1) * items_per_page
    end = start + items_per_page
    paginated_data = sorted_data.iloc[start:end]

    # Generate page numbers for display
    pagination_numbers = list(range(1, total_pages + 1))

    # Pass all necessary variables to the template
    return render_template(
        'themes.html',
        themes_data=paginated_data.to_dict(orient='records'),
        page=page,  # Current page
        total_pages=total_pages  # Total number of pages
    )

# Helper functions
def safe_str(value):
    return str(value).strip() if pd.notna(value) else ""

def format_authors(authors):
    if not authors or pd.isna(authors):
        return "Author unknown"
    author_list = re.split(r';|,', authors)
    formatted_authors = []
    for author in author_list:
        parts = author.strip().split()
        if len(parts) >= 2:
            last_name = parts[-1]
            initials = "".join([f"{name[0]}." for name in parts[:-1]])
            formatted_authors.append(f"{last_name}, {initials}")
        else:
            formatted_authors.append(author.strip())
    return f"{', '.join(formatted_authors[:-1])}, & {formatted_authors[-1]}" if len(formatted_authors) > 1 else formatted_authors[0]

def format_apa(row):
    ref_no = safe_str(row.get("Reference_no"))
    url = safe_str(row.get("URL"))
    authors = format_authors(safe_str(row.get("Main_Author")))
    title = safe_str(row.get("Title"))
    journal_name = safe_str(row.get("Journal Name"))
    pub_date = safe_str(row.get("Publication Date"))
    year_match = re.search(r'\d{4}', pub_date)
    year = year_match.group() if year_match else "n.d."

    # Format the reference_no as a clickable link
    ref_link = f"<a href='{url}' target='_blank' class='text-blue-500 hover:underline'>{ref_no}</a>"

    # Build the full citation
    citation = f"({ref_link}) {authors}. ({year}). {title}."
    if journal_name:
        citation += f" {journal_name}."
    return citation

@app.route('/references')
def references():
    """
    Route to display references with pagination and clickable links.
    Supports view_all parameter to show all references at once.
    """
    view_all = request.args.get('view_all', 'false').lower() == 'true'
    
    if view_all:
        # Show all references without pagination
        # Create references with original reference numbers for view all
        formatted_references = []
        for _, row in data.iterrows():
            ref_no = safe_str(row.get("Reference_no"))
            url = safe_str(row.get("URL"))
            authors = format_authors(safe_str(row.get("Main_Author")))
            title = safe_str(row.get("Title"))
            journal_name = safe_str(row.get("Journal Name"))
            pub_date = safe_str(row.get("Publication Date"))
            year_match = re.search(r'\d{4}', pub_date)
            year = year_match.group() if year_match else "n.d."
            
            # Build citation with original reference number
            citation_data = {
                'ref_no': ref_no,
                'authors': authors,
                'year': year,
                'title': title,
                'journal': journal_name,
                'url': url
            }
            formatted_references.append(citation_data)
        
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        return render_template(
            'references_all.html',
            references=formatted_references,
            total_count=len(data),
            current_date=current_time
        )
    else:
        # Normal pagination
        items_per_page = 20
        page = request.args.get('page', default=1, type=int)
        total_items = len(data)
        total_pages = math.ceil(total_items / items_per_page)

        # Validate page
        if page < 1:
            page = 1
        elif page > total_pages:
            page = total_pages

        start = (page - 1) * items_per_page
        end = start + items_per_page
        paginated_data = data.iloc[start:end]

        # Format the references
        formatted_references = [format_apa(row) for _, row in paginated_data.iterrows()]

        return render_template(
            'references.html',
            references=formatted_references,
            page=page,
            total_pages=total_pages
        )


@app.route('/download/<format_type>')
def download_references(format_type):
    """
    Route to download all references in different formats (PDF, Word, TXT).
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from io import BytesIO
        
        # Generate formatted references with URLs for downloads
        all_references = []
        all_references_with_urls = []
        
        for _, row in data.iterrows():
            formatted_ref = format_apa(row)
            # Clean HTML tags for TXT format
            clean_ref = re.sub(r'<[^>]+>', '', formatted_ref)
            all_references.append(clean_ref)
            
            # Keep reference with URL info for Word/PDF
            ref_data = {
                'clean_text': clean_ref,
                'url': safe_str(row.get("URL")),
                'ref_no': safe_str(row.get("Reference_no")),
                'authors': format_authors(safe_str(row.get("Main_Author"))),
                'title': safe_str(row.get("Title")),
                'journal': safe_str(row.get("Journal Name")),
                'year': re.search(r'\d{4}', safe_str(row.get("Publication Date"))).group() if re.search(r'\d{4}', safe_str(row.get("Publication Date"))) else "n.d."
            }
            all_references_with_urls.append(ref_data)
        
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        if format_type == 'txt':
            # Generate TXT file with URLs
            content = f"Recommender Systems Systematic Review - References\n"
            content += f"Generated on: {current_date}\n"
            content += f"Total References: {len(all_references_with_urls)}\n"
            content += "=" * 80 + "\n\n"
            
            for i, ref_data in enumerate(all_references_with_urls, 1):
                # Build reference with URL for TXT format
                ref_text = f"({ref_data['ref_no']}) {ref_data['authors']} ({ref_data['year']}). {ref_data['title']}."
                if ref_data['journal']:
                    ref_text += f" {ref_data['journal']}."
                if ref_data['url']:
                    ref_text += f"\n    View Paper: {ref_data['url']}"
                
                content += f"{i:3d}. {ref_text}\n\n"
            
            output = BytesIO()
            output.write(content.encode('utf-8'))
            output.seek(0)
            
            return send_file(
                output,
                as_attachment=True,
                download_name=f"RS_References_{current_date}.txt",
                mimetype="text/plain"
            )
            
        elif format_type == 'word':
            # Generate proper Word document using python-docx
            doc = Document()
            
            # Set document title
            title = doc.add_heading('Recommender Systems Systematic Review', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            subtitle = doc.add_heading('Complete References (APA Style)', level=1)
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            meta_run = meta_para.add_run(f"Generated on: {current_date}\nTotal References: {len(all_references)}")
            meta_run.italic = True
            
            doc.add_paragraph()  # Add space
            
            # Add references with clickable links (simplified approach)
            for i, ref_data in enumerate(all_references_with_urls, 1):
                ref_para = doc.add_paragraph()
                
                # Add reference number in bold
                num_run = ref_para.add_run(f"{i:3d}. ")
                num_run.bold = True
                
                # Add reference text
                ref_para.add_run(f"{ref_data['authors']} ({ref_data['year']}). {ref_data['title']}.")
                if ref_data['journal']:
                    ref_para.add_run(f" {ref_data['journal']}.")
                
                # Add URL as clickable link if available
                if ref_data['url']:
                    ref_para.add_run(" ")
                    # Add URL as plain text since hyperlinks in python-docx are complex
                    url_run = ref_para.add_run(f"[View Paper: {ref_data['url']}]")
                    from docx.shared import RGBColor
                    url_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
                    url_run.italic = True
                
                # Add hanging indent for APA style
                ref_para.paragraph_format.left_indent = Inches(0.5)
                ref_para.paragraph_format.first_line_indent = Inches(-0.5)
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return send_file(
                output,
                as_attachment=True,
                download_name=f"RS_References_{current_date}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        elif format_type == 'pdf':
            # Generate proper PDF using ReportLab
            output = BytesIO()
            doc = SimpleDocTemplate(output, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=12,
                alignment=1,  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                alignment=1,  # Center alignment
            )
            
            meta_style = ParagraphStyle(
                'MetaInfo',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=20,
                alignment=1,  # Center alignment
                textColor=colors.grey,
            )
            
            # Custom reference style with hanging indent
            ref_style = ParagraphStyle(
                'Reference',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leftIndent=36,
                firstLineIndent=-36,
                alignment=0,  # Left alignment
                leading=14,
            )
            
            # Build content
            content = []
            
            # Add title and metadata
            content.append(Paragraph("Recommender Systems Systematic Review", title_style))
            content.append(Paragraph("Complete References (APA Style)", subtitle_style))
            content.append(Paragraph(f"Generated on: {current_date}<br/>Total References: {len(all_references)}", meta_style))
            content.append(Spacer(1, 20))
            
            # Add references with clickable links
            for i, ref_data in enumerate(all_references_with_urls, 1):
                if ref_data['url']:
                    # Create reference with clickable number
                    ref_text = f'<b><a href="{ref_data["url"]}" color="blue">{i:3d}.</a></b> {ref_data["authors"]} ({ref_data["year"]}). <i>{ref_data["title"]}</i>.'
                    if ref_data['journal']:
                        ref_text += f' <i>{ref_data["journal"]}</i>.'
                    ref_text += f' <a href="{ref_data["url"]}" color="blue">[View Paper]</a>'
                else:
                    # No URL available, use clean reference
                    ref_text = f"<b>{i:3d}.</b> {ref_data['clean_text']}"
                
                content.append(Paragraph(ref_text, ref_style))
            
            # Build PDF
            doc.build(content)
            output.seek(0)
            
            return send_file(
                output,
                as_attachment=True,
                download_name=f"RS_References_{current_date}.pdf",
                mimetype="application/pdf"
            )
            
        else:
            return jsonify({"error": "Invalid format. Use 'pdf', 'word', or 'txt'"}), 400
            
    except Exception as e:
        print(f"Download error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Download failed: {str(e)}"}), 500

@app.route('/download/<classified_url>')
def download_excel(classified_url):
    """
    Route to download an Excel file for a specific 'Classified' value.
    """
    # Decode the classified_url and map it to the original Classified value
    classified_value = url_to_classified.get(unquote(classified_url))
    print(f"Received classified_url: {classified_url}")  # Debug log
    print(f"Mapped classified_value: {classified_value}")  # Debug log
    
    if not classified_value or classified_value not in grouped_data.groups:
        return jsonify({"error": "Invalid classified value."}), 404

    # Get filtered data for the classified value
    filtered_data = grouped_data.get_group(classified_value)

    # Save the filtered data to an Excel file in memory
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        filtered_data.to_excel(writer, index=False, sheet_name="Data")
    output.seek(0)

    # Generate a safe filename
    safe_filename = f"{classified_url}.xlsx"

    return send_file(
        output,
        as_attachment=True,
        download_name=safe_filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# Run the app
if __name__ == "__main__":
    # Bind the app to the PORT environment variable
    port = int(os.environ.get("PORT", 8080))
    # Only debug in local development (not in production)
    debug_mode = not (os.environ.get("RENDER") or os.environ.get("AWS_DEPLOYMENT"))
    app.run(debug=debug_mode, host="0.0.0.0", port=port)